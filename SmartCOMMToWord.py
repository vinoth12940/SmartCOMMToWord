import os
import base64
from io import BytesIO
from PIL import Image
import xml.etree.ElementTree as ET
import PySimpleGUI as sg
import docx
from docx.shared import Cm

# Define the input and output directories
xmlFileFolder = ''
newWordDocFolder = ''

# Define the layout of the GUI
description = sg.Text('XML to Word Converter', font=('Helvetica', 20), justification='center')
prerequisites = sg.Text('Prerequisites:\n\n1. Python 3.x\n2. PySimpleGUI\n3. Pillow\n4. python-docx\n\n', font=('Helvetica', 14), justification='center')
instructions = sg.Text('Instructions:\n\n1. Select the XML file folder.\n2. Select the new Word document folder.\n3. Click the "Convert" button to convert the XML files to Word documents.\n\n', font=('Helvetica', 14), justification='center')
description_box = sg.Column([[description]], element_justification='center')
prerequisites_box = sg.Column([[prerequisites]], element_justification='center')
instructions_box = sg.Column([[instructions]], element_justification='center')
layout = [
    [description_box],
    [sg.Frame('Prerequisites', [[prerequisites_box]], element_justification='center')],
    [sg.Frame('Instructions', [[instructions_box]], element_justification='center')],
    [sg.Text('XML File Folder:'), sg.Input(xmlFileFolder, key='xmlFileFolder'), sg.FolderBrowse()],
    [sg.Text('New Word Doc Folder:'), sg.Input(newWordDocFolder, key='newWordDocFolder'), sg.FolderBrowse()],
    [sg.Button('Convert'), sg.Button('Exit')]
]

# Create the GUI window
window = sg.Window('XML to Word Converter', layout, size=(900, 500))

# Define a function to add images to the document
def add_image(doc, image_data, width, height, format, checksum):
    # Decode the Base64-encoded image data
    image_bytes = base64.b64decode(image_data)
    image_stream = BytesIO(image_bytes)
    img = Image.open(image_stream)
    img.save("temp_image." + format)
    doc.add_picture("temp_image." + format, width=docx.shared.Inches(width/96), height=docx.shared.Inches(height/96))  # You may need to adjust the width and height

# Define a function to convert the XML files to Word documents
def convert_files(xmlFileFolder, newWordDocFolder):
    # Create the output directory if it does not exist
    if not os.path.exists(newWordDocFolder):
        os.makedirs(newWordDocFolder)

    # Check if there are any XML files in the input directory
    xml_files = [f for f in os.listdir(xmlFileFolder) if f.endswith('.xml')]
    if not xml_files:
        raise Exception('No XML files found in the input directory')

    # Loop through all the XML files in the input directory
    for xml_file in xml_files:
        # Open the XML file
        tree = ET.parse(os.path.join(xmlFileFolder, xml_file))
        root = tree.getroot()

        # Create a new Word document
        doc = docx.Document()

        # Define a custom style for the hyperlink
        hyperlink_style = doc.styles.add_style('Hyperlink', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        hyperlink_style.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
        hyperlink_style.font.underline = True

        # Loop through the XML elements and add them to the document
        for elem in root.iter():
            if elem.text:
                if elem.tag == 'p':
                    if 'style' in elem.attrib and 'Heading' in elem.attrib['style']:
                        doc.add_heading(elem.text, int(elem.attrib['style'][-1]))
                    else:
                        doc.add_paragraph(elem.text)
                elif elem.tag == 'ulist':
                    for li in elem.iter('listitem'):
                        doc.add_paragraph(li.text, style='List Bullet')
                elif elem.tag == 'hyperlink':
                    ref = elem.find('ref')
                    if ref is not None:
                        url = ref.find('string').text
                        display = elem.find('display/string').text
                        doc.add_paragraph(display, style='Hyperlink').hyperlink = url
                elif elem.tag == 'image':
                    width = int(elem.attrib['width'][:-2])
                    height = int(elem.attrib['height'][:-2])
                    format = elem.attrib['format']
                    checksum = elem.attrib['checksum']
                    base64_data = elem.text
                    add_image(doc, base64_data, width, height, format, checksum)
                elif elem.tag == 'table':
                    table_data = []
                    for row_elem in elem.findall(".//row"):
                        row_data = []
                        for cell_elem in row_elem.findall(".//cell"):
                            cell_text = cell_elem.find(".//p").text
                            row_data.append(cell_text)
                        table_data.append(row_data)

                    # Add the table to the document
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    for i in range(len(table_data)):
                        for j in range(len(table_data[i])):
                            table.cell(i, j).text = table_data[i][j]
                    table.style = 'Table Grid'
                elif elem.tag == 'string':
                    doc.add_paragraph(elem.text)
                elif elem.tag == 'style':
                    doc.add_paragraph(elem.text)                    

        # Save the Word document
        doc.save(os.path.join(newWordDocFolder, os.path.splitext(xml_file)[0] + '.docx'))

# Event loop for the GUI
while True:
    event, values = window.read()
    if event == sg.WIN_CLOSED or event == 'Exit':
        break
    elif event == 'Convert':
        xmlFileFolder = values['xmlFileFolder']
        newWordDocFolder = values['newWordDocFolder']
        try:
            convert_files(xmlFileFolder, newWordDocFolder)
            sg.popup('Conversion complete!')
        except Exception as e:
            sg.popup_error(f'Error: {str(e)}')

# Close the GUI window
window.close()