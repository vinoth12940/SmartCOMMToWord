import os
import docx
from xml.etree import ElementTree
import base64
from io import BytesIO
from PIL import Image
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Cm

# Define the input and output directories
xmlFileFolder = '/Users/vinothrajalingam/Desktop/Python/xmlFileFolder'
newWordDocFolder = '/Users/vinothrajalingam/Desktop/Python/newWordDocFolder'

# Create the output directory if it does not exist
if not os.path.exists(newWordDocFolder):
    os.makedirs(newWordDocFolder)

# Define a function to add images to the document
def add_image(doc, image_data, width, height, format, checksum):
    # Decode the Base64-encoded image data
    image_bytes = base64.b64decode(image_data)
    image_stream = BytesIO(image_bytes)
    img = Image.open(image_stream)
    img.save("temp_image." + format)
    doc.add_picture("temp_image." + format, width=docx.shared.Inches(width/96), height=docx.shared.Inches(height/96))  # You may need to adjust the width and height

# Loop through all the XML files in the input directory
for xml_file in os.listdir(xmlFileFolder):
    if xml_file.endswith('.xml'):
        # Open the XML file
        tree = ElementTree.parse(os.path.join(xmlFileFolder, xml_file))
        root = tree.getroot()

        # Create a new Word document
        doc = docx.Document()

        # Define custom styles for the document
        existing_styles = [s.name for s in doc.styles]
        for style in root.findall('.//style'):
            name = style.attrib.get('name')
            if name in existing_styles:
                doc.styles[name].delete()
            if name and name not in existing_styles:
                try:
                    doc.styles.add_style(name, docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
                except ValueError:
                    pass

        # Define the 'Hyperlink' style
        hyperlink_style = doc.styles.add_style('Hyperlink', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        hyperlink_style.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
        hyperlink_style.font.underline = True

        # Loop through the XML elements and add them to the document
        for elem in root.iter():
            if elem.text:
                if elem.tag == 'p':
                    if 'style' in elem.attrib and 'Heading' in elem.attrib['style']:
                        doc.add_heading(elem.text, int(elem.attrib['style'][-1]))
                    else:
                        text = ''
                        for child in elem.iter():
                            if child.text:
                                text += child.text
                        doc.add_paragraph(text)
                elif elem.tag == 'ulist':
                    for li in elem.iter('listitem'):
                        doc.add_paragraph(li.text, style='List Bullet')
                elif elem.tag == 'hyperlink':
                    ref = elem.find('ref')
                    if ref is not None:
                        url = ref.find('string').text
                        display = elem.find('display/string').text
                        doc.add_paragraph(display, style='Hyperlink').hyperlink = url
                elif elem.tag == 'image':
                    width = int(elem.attrib['width'][:-2])
                    height = int(elem.attrib['height'][:-2])
                    format = elem.attrib['format']
                    checksum = elem.attrib['checksum']
                    base64_data = elem.text
                    add_image(doc, base64_data, width, height, format, checksum)
                elif elem.tag == 'table':
                    table_data = []
                    for row_elem in elem.findall(".//row"):
                        row_data = []
                        for cell_elem in row_elem.findall(".//cell"):
                            cell_text = cell_elem.find(".//p").text
                            row_data.append(cell_text)
                        table_data.append(row_data)

                    # Add the table to the document
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    for i in range(len(table_data)):
                        for j in range(len(table_data[i])):
                            table.cell(i, j).text = table_data[i][j]
                    table.style = 'Table Grid'

        # Save the Word document
        doc.save(os.path.join(newWordDocFolder, os.path.splitext(xml_file)[0] + '.docx'))